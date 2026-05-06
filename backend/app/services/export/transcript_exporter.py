import os
import uuid
from pathlib import Path

from app.models.meeting import Meeting
from app.models.segment import TranscriptSegment


def _fmt_ms(ms: int) -> str:
    s = ms // 1000
    m, sec = divmod(s, 60)
    return f"{m:02d}:{sec:02d}"


def _resolve_speaker(label: str | None, speaker_map: dict[str, str] | None) -> str:
    if not label:
        return ""
    if speaker_map and label in speaker_map:
        return speaker_map[label]
    return label


def _build_md(meeting: Meeting, segments: list[TranscriptSegment]) -> str:
    spk_map: dict[str, str] | None = getattr(meeting, "speaker_map", None)
    lines: list[str] = []
    lines.append(f"# {meeting.title}")
    lines.append("")
    if meeting.meeting_time:
        lines.append(f"**时间：** {meeting.meeting_time.strftime('%Y-%m-%d %H:%M')}")
    if meeting.participants:
        lines.append(f"**参会人：** {', '.join(meeting.participants)}")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## 转写记录")
    lines.append("")
    for seg in segments:
        text = seg.corrected_text or seg.raw_text or ""
        name = _resolve_speaker(seg.speaker_label, spk_map)
        speaker = f"[{name}] " if name else ""
        time_tag = f"`{_fmt_ms(seg.start_ms)}–{_fmt_ms(seg.end_ms)}`"
        lines.append(f"{time_tag} {speaker}{text}")
        lines.append("")
    return "\n".join(lines)


def export_md(
    meeting: Meeting,
    segments: list[TranscriptSegment],
    export_dir: str,
) -> str:
    content = _build_md(meeting, segments)
    filename = f"transcript_{meeting.id}_{uuid.uuid4().hex[:8]}.md"
    path = Path(export_dir) / filename
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return str(path)


def export_docx(
    meeting: Meeting,
    segments: list[TranscriptSegment],
    export_dir: str,
) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title_para = doc.add_heading(meeting.title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if meeting.meeting_time:
        doc.add_paragraph(f"时间：{meeting.meeting_time.strftime('%Y-%m-%d %H:%M')}")
    if meeting.participants:
        doc.add_paragraph(f"参会人：{', '.join(meeting.participants)}")

    doc.add_paragraph()
    doc.add_heading("转写记录", level=2)

    spk_map: dict[str, str] | None = getattr(meeting, "speaker_map", None)
    for seg in segments:
        text = seg.corrected_text or seg.raw_text or ""
        name = _resolve_speaker(seg.speaker_label, spk_map)
        speaker = f"[{name}] " if name else ""
        time_tag = f"{_fmt_ms(seg.start_ms)}–{_fmt_ms(seg.end_ms)}"
        p = doc.add_paragraph()
        run_time = p.add_run(f"{time_tag}  ")
        run_time.font.size = Pt(9)
        run_time.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        if speaker:
            run_spk = p.add_run(speaker)
            run_spk.bold = True
        p.add_run(text)

    filename = f"transcript_{meeting.id}_{uuid.uuid4().hex[:8]}.docx"
    path = Path(export_dir) / filename
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return str(path)
