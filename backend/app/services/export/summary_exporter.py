import uuid
from pathlib import Path

from app.models.meeting import Meeting
from app.models.summary import MeetingSummary


def export_md(
    meeting: Meeting,
    summary: MeetingSummary,
    export_dir: str,
) -> str:
    title = summary.title or meeting.title
    lines: list[str] = []
    lines.append(f"# {title}")
    lines.append("")
    if meeting.meeting_time:
        lines.append(f"**时间：** {meeting.meeting_time.strftime('%Y-%m-%d %H:%M')}")
    if meeting.participants:
        lines.append(f"**参会人：** {', '.join(meeting.participants)}")
    if summary.llm_model:
        lines.append(f"**生成模型：** {summary.llm_model}")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append(summary.content_md or "")
    content = "\n".join(lines)

    filename = f"summary_{summary.id}_{uuid.uuid4().hex[:8]}.md"
    path = Path(export_dir) / filename
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    return str(path)


def export_docx(
    meeting: Meeting,
    summary: MeetingSummary,
    export_dir: str,
) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = summary.title or meeting.title

    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_lines = []
    if meeting.meeting_time:
        meta_lines.append(f"时间：{meeting.meeting_time.strftime('%Y-%m-%d %H:%M')}")
    if meeting.participants:
        meta_lines.append(f"参会人：{', '.join(meeting.participants)}")
    if summary.llm_model:
        meta_lines.append(f"生成模型：{summary.llm_model}")
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    content = summary.content_md or ""
    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("| ") and stripped.endswith(" |"):
            # Simple table row — add as plain paragraph
            doc.add_paragraph(stripped)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)
        else:
            doc.add_paragraph()

    filename = f"summary_{summary.id}_{uuid.uuid4().hex[:8]}.docx"
    path = Path(export_dir) / filename
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return str(path)
